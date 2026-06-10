"""Walk submission/ and convert every Markdown file to a sibling .docx.

Uses python-docx to render headings, paragraphs, bullet lists, ordered lists,
code blocks, tables, blockquotes, and inline bold/italic/code formatting.

This is a minimal Markdown subset — no images, no nested lists. It covers the
AFCN submission corpus cleanly without dragging in pandoc.
"""
from __future__ import annotations
import re, sys, html
from pathlib import Path
sys.stdout.reconfigure(encoding="utf-8")

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("pip install python-docx")

ROOT = Path(__file__).resolve().parent

# ── inline-formatting parser ──────────────────────────────────────────
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*)"           # bold
    r"|(\*[^*]+\*)"              # italic (asterisk)
    r"|(__[^_]+__)"              # bold (underscore)
    r"|(_[^_\s][^_]*_)"          # italic (underscore)
    r"|(`[^`]+`)"                # inline code
    r"|(\[[^\]]+\]\([^)]+\))"    # link [text](url)
)


def add_inline_runs(par, text: str) -> None:
    """Append runs to a paragraph, respecting bold / italic / code / link."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            par.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**") or token.startswith("__"):
            par.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = par.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            # Light grey shading via XML
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F0F0F0")
            run._element.get_or_add_rPr().append(shd)
        elif token.startswith("["):
            # [label](url) → render label with italic blue
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if mm:
                run = par.add_run(mm.group(1))
                run.font.color.rgb = RGBColor(0x1C, 0x5C, 0x84)  # Nr. 51
                run.font.underline = True
        elif token.startswith("*") or token.startswith("_"):
            par.add_run(token[1:-1]).italic = True
        else:
            par.add_run(token)
        pos = m.end()
    if pos < len(text):
        par.add_run(text[pos:])


# ── block-level converter ─────────────────────────────────────────────
def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    # Base styling: serif body, sans headings
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for s in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
        h = doc.styles[s]
        h.font.name = "Calibri"
        h.font.color.rgb = RGBColor(0x39, 0x3F, 0x44)  # Nr. 49 ink

    lines = text.splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_buf: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table():
        if not table_rows:
            return
        # First row is header, second is separator (drop), rest is body
        body = [r for r in table_rows[2:]] if len(table_rows) >= 2 else table_rows
        header = table_rows[0]
        cols = len(header)
        tbl = doc.add_table(rows=1 + len(body), cols=cols)
        tbl.style = "Light Grid Accent 1"
        for c, txt in enumerate(header):
            cell = tbl.rows[0].cells[c]
            p = cell.paragraphs[0]
            run = p.add_run(txt.strip())
            run.bold = True
        for r, row in enumerate(body, start=1):
            for c in range(cols):
                cell = tbl.rows[r].cells[c]
                p = cell.paragraphs[0]
                val = row[c].strip() if c < len(row) else ""
                add_inline_runs(p, val)
        doc.add_paragraph()  # spacer after table

    def flush_code():
        if not code_buf:
            return
        p = doc.add_paragraph()
        # Light grey background shading on the whole paragraph
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F4F4F4")
        pPr.append(shd)
        run = p.add_run("\n".join(code_buf))
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x39, 0x3F, 0x44)

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            if in_code:
                flush_code()
                code_buf.clear()
                in_code = False
            else:
                in_code = True
                code_lang = line[3:].strip()
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Table (GitHub style)
        if "|" in line and (i + 1 < len(lines)) and re.match(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[i + 1] or ""):
            # Collect table rows
            in_table = True
            table_rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                row = [c for c in re.split(r"(?<!\\)\|", lines[i].strip().strip("|"))]
                table_rows.append(row)
                i += 1
            flush_table()
            in_table = False
            table_rows = []
            continue

        # Horizontal rule
        if re.match(r"^\s*---+\s*$", line):
            doc.add_paragraph().add_run("─" * 60).font.color.rgb = RGBColor(0xA4, 0xB1, 0xBC)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i].lstrip("> ").rstrip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            run = p.add_run("  ")
            # Vertical bar (use border-left via XML)
            pPr = p._element.get_or_add_pPr()
            bdr = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "12")
            left.set(qn("w:color"), "7D9CAF")  # Nr. 57
            bdr.append(left)
            pPr.append(bdr)
            add_inline_runs(p, " ".join(quote_lines))
            run = p.add_run("")
            for r in p.runs:
                r.italic = True
            continue

        # Bullet list
        if re.match(r"^\s*[-*+]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*+]\s+", lines[i] or ""):
                m = re.match(r"^\s*[-*+]\s+(.*)$", lines[i])
                p = doc.add_paragraph(style="List Bullet")
                add_inline_runs(p, m.group(1))
                i += 1
            continue

        # Ordered list
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i] or ""):
                m = re.match(r"^\s*\d+\.\s+(.*)$", lines[i])
                p = doc.add_paragraph(style="List Number")
                add_inline_runs(p, m.group(1))
                i += 1
            continue

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (collect consecutive non-blank, non-special lines)
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if (not nxt.strip()
                or nxt.startswith(("#", "```", ">", "---"))
                or re.match(r"^\s*([-*+]|\d+\.)\s+", nxt)
                or ("|" in nxt and (i + 1 < len(lines)) and re.match(r"^\s*\|?\s*:?-+:?", lines[i + 1] or ""))):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(line.rstrip() for line in para_lines))

    # Flush trailing code if any
    if in_code:
        flush_code()

    doc.save(docx_path)


# ── walker ────────────────────────────────────────────────────────────
def main() -> None:
    # Skip the legacy _docx shadow tree if it exists
    md_files = [p for p in sorted(ROOT.rglob("*.md"))
                if "_docx" not in p.parts]
    print(f"Found {len(md_files)} .md files under {ROOT}")
    for md in md_files:
        rel = md.relative_to(ROOT)
        # Write .docx alongside the source .md (not in a shadow tree)
        docx_path = md.with_suffix(".docx")
        try:
            convert_md_to_docx(md, docx_path)
            print(f"  + {rel.with_suffix('.docx')}")
        except Exception as e:
            print(f"  ! {rel}: {type(e).__name__}: {e}")
    print(f"\nDone. .docx files written alongside their .md sources.")


if __name__ == "__main__":
    main()
